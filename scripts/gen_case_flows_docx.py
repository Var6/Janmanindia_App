#!/usr/bin/env python3
"""Generate the editable Word doc of case types & step flows for the funder meeting.

Updated to the FOUR-FLOW model:
  1. Criminal           (FIR ... Verdict, + optional Bail track)
  2. Family Court       (8-step)
  3. Other / Civil      (6-step "Other Cases")
  4. Writ & Appeals     (7-step petition flow)
Case types are the real eCourts catalog (High Court + District), each tagged
with the flow it follows.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/rishabhranjan/Desktop/Janmanindia_App/Case-Flows.docx"

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
RED = RGBColor(0xB0, 0x2A, 0x2A)
BLUE = RGBColor(0x1A, 0x4D, 0x8F)
GREEN = RGBColor(0x1E, 0x7A, 0x3C)
PURPLE = RGBColor(0x6B, 0x2A, 0x8F)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def shade(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hex_color)
    tcPr.append(sh)


def set_cell_text(cell, text, bold=False, color=None, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def heading(text, size=16, color=NAVY, space_before=14, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def para(text, italic=False, color=None, size=11, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def change_box(label="✍️  CHANGE / NOTES:"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.cell(0, 0)
    shade(cell, "FFF7E0")
    p = cell.paragraphs[0]
    r = p.add_run(label + "  ")
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0x99, 0x66, 0x00)
    cell.add_paragraph("")
    cell.add_paragraph("")
    doc.add_paragraph("")


def make_table(headers, rows, widths=None, header_fill="1F3A5F"):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10)
        shade(hdr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=10)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def steps_table(steps, color_hex):
    rows = []
    for i, (label, editable) in enumerate(steps, start=1):
        rows.append([str(i), label, editable])
    make_table(["#", "Step (what the user sees)", "Editable by staff?"], rows,
               widths=[0.5, 4.1, 1.9], header_fill=color_hex)


# ── TITLE ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("Case Types & Step Flows")
r.bold = True
r.font.size = Pt(24)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Editable Reference  ·  Janman India Case Management")
r.font.size = Pt(12)
r.font.color.rgb = GREY

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run("Prepared for funder review  ·  16 June 2026")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = GREY

doc.add_paragraph("")

ht = doc.add_table(rows=1, cols=1)
ht.style = "Table Grid"
c = ht.cell(0, 0)
shade(c, "EAF1FB")
c.text = ""
p = c.paragraphs[0]
r = p.add_run("How to use this document\n")
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = NAVY
for line in [
    "• Each step shows what the user sees. Every step can now be clicked to mark it done.",
    "• Write edits in the yellow CHANGE / NOTES boxes — rename, add, remove, or reorder a step.",
    "• Each case type is tagged with the flow it follows. Note any type that should switch flows.",
]:
    pp = c.add_paragraph(line)
    pp.runs[0].font.size = Pt(10)

doc.add_paragraph("")

# ── BIG PICTURE ──────────────────────────────────────────────────────────────
heading("The Big Picture (read this first)", size=18)
para("Every case carries two labels:")
p = doc.add_paragraph(style="List Number")
p.add_run("Case Type").bold = True
p.add_run(" — what kind of case it is (the eCourts catalog: FIR, Title Suit, Writ, Divorce…). 150+ types.")
p = doc.add_paragraph(style="List Number")
p.add_run("Flow").bold = True
p.add_run(" — which set of steps it follows. There are FOUR flows. Each case type maps to one.")

para("")
make_table(
    ["Flow", "Used for"],
    [
        ["🔴 Criminal", "FIR, sessions, bail, POCSO/NDPS/SC-ST & other criminal special acts"],
        ["🟢 Family Court", "Matrimonial, divorce, guardianship, maintenance, special marriage"],
        ["🔵 Other / Civil", "Suits — title, money, partition, eviction, succession, execution"],
        ["🟣 Writ & Appeals", "Writ petitions, appeals, revisions, references, company/tax petitions"],
    ],
    widths=[1.7, 4.3],
)
para("")
para("To move a case type to a different flow is a one-line change. Note it in the boxes below.",
     italic=True, color=GREEN)

doc.add_page_break()

# ── PART 1 — THE FOUR FLOWS ──────────────────────────────────────────────────
heading("Part 1 — The Four Step-Flows", size=18)

# Flow A — Criminal
heading("🔴 Flow 1 — Criminal", size=14, color=RED)
para("Used for FIR, sessions trials, magistrate complaints and criminal special acts.")
steps_table([
    ("Case Filed", "Automatic"),
    ("FIR Filed", "Yes — click to mark"),
    ("Investigation", "Yes — click to mark"),
    ("Chargesheet Filed", "Yes — click to mark"),
    ("Cognizance", "Yes — click to mark"),
    ("Appearance of Accused", "Yes — click to mark"),
    ("Framing of Charges", "Yes — click to mark"),
    ("Prosecution Witnesses", "Yes — click to mark"),
    ("Defence Witnesses", "Yes — click to mark"),
    ("Examination of Accused (Sec 313 BNSS)", "Yes — click to mark"),
    ("Arguments", "Yes — click to mark"),
    ("Verdict → Acquittal / Conviction → Sentencing → Closed", "Yes — click to mark"),
], "B02A2A")
para("")
para("Optional Bail track (opt-in): Case Filed → Bail Application Filed → Bail Hearing → "
     "Bail Granted / Bail Rejected.", italic=True, color=GREY, size=10)
change_box()

doc.add_page_break()

# Flow B — Family Court
heading("🟢 Flow 2 — Family Court", size=14, color=GREEN)
para("Your definition: matrimonial, divorce, guardianship, maintenance, special marriage.")
steps_table([
    ("Case Filed", "Automatic"),
    ("Admission", "Yes — click to mark"),
    ("Appearance of Opposite Party", "Yes — click to mark"),
    ("Written Statement", "Yes — click to mark"),
    ("Framing of Issues", "Yes — click to mark"),
    ("Witness", "Yes — click to mark"),
    ("Arguments", "Yes — click to mark"),
    ("Final Outcome — Allowed / Disposed", "Yes — click to mark"),
], "1E7A3C")
change_box()

# Flow C — Other / Civil
heading("🔵 Flow 3 — Other / Civil", size=14, color=BLUE)
para("Your 'Other Cases' definition: general civil suits and applications.")
steps_table([
    ("Case Filed", "Automatic"),
    ("Hearing on Admission", "Yes — click to mark"),
    ("Written Statement / Counter Affidavit", "Yes — click to mark"),
    ("Rejoinder", "Yes — click to mark"),
    ("Final Hearing", "Yes — click to mark"),
    ("Outcome — Allowed / Disposed", "Yes — click to mark"),
], "1A4D8F")
change_box()

doc.add_page_break()

# Flow D — Writ & Appeals
heading("🟣 Flow 4 — Writ & Appeals", size=14, color=PURPLE)
para("Writ petitions, appeals, revisions, references, company / tax petitions (the existing High Court flow).")
steps_table([
    ("Case Filed", "Automatic"),
    ("Petition Filed", "Yes — click to mark"),
    ("Supporting Affidavit", "Yes — click to mark"),
    ("Admission", "Yes — click to mark"),
    ("Counter Affidavit", "Yes — click to mark"),
    ("Rejoinder", "Yes — click to mark"),
    ("Plea Close", "Yes — click to mark"),
    ("Inducement / Judgment", "Yes — click to mark"),
], "6B2A8F")
change_box()

# Status
heading("Overall Case Status (separate from the steps)", size=13)
make_table(
    ["Status", "Meaning"],
    [
        ["Pending", "Being heard"],
        ["Disposal", "Disposed of by the court"],
        ["Withdrawn", "Withdrawn by the party"],
    ],
    widths=[1.6, 4.4],
)
para("")
change_box()

doc.add_page_break()

# ── PART 2 — CASE TYPES (real eCourts catalog) ───────────────────────────────
heading("Part 2 — Case Types (eCourts catalog)", size=18)
para("These are the real case types from the eCourts dropdowns, grouped by court and tagged with "
     "the flow each follows. 🔴 Criminal · 🟢 Family · 🔵 Civil · 🟣 Writ/Appeals.")
para("")

heading("High Court — Criminal", size=12, color=NAVY, space_before=8)
make_table(
    ["Code", "Case type", "Flow"],
    [
        ["CR. APP (DB/SJ/US)", "Criminal Appeal (Division/Single Bench, U/S)", "🔴"],
        ["CR. APP(341)", "Cr. Appeal u/s 341 CrPC", "🔴"],
        ["CR. MISC.", "Criminal Miscellaneous", "🔴"],
        ["OR. CR. MISC", "Original Criminal Miscellaneous (+ DB)", "🔴"],
        ["CR. REV.", "Criminal Revision", "🔴"],
        ["CR. REF. / D. REF.", "Criminal Reference / Death Reference", "🔴"],
        ["CR. WJC", "Criminal Writ Jurisdiction Case", "🟣"],
        ["G. APP (DB/SJ)", "Govt. Appeal (Division/Single Bench)", "🔴"],
        ["SLA", "Special Leave Application", "🟣"],
    ],
    widths=[1.7, 3.6, 0.7],
)

heading("High Court — Civil", size=12, color=NAVY, space_before=10)
make_table(
    ["Code", "Case type", "Flow"],
    [
        ["FA / SA", "First Appeal / Second Appeal", "🟣"],
        ["MA", "Miscellaneous Appeal", "🟣"],
        ["L.P.A", "Letters Patent Appeal", "🟣"],
        ["C.R. / C. REV. / C.REF.", "Civil Revision / Review / Reference", "🟣"],
        ["CWJC / MJC", "Civil Writ Jurisdiction / Misc. Jurisdiction Case", "🟣"],
        ["COMP. APP / PET / APPLIC", "Company Appeal / Petition / Application", "🟣"],
        ["COMMERCIAL APP", "Commercial Appeal", "🟣"],
        ["CRT. / COMPAN. / SCA", "Certificate / Compensation / Supreme Court Appeal", "🟣"],
        ["E.P.", "Election Petition", "🟣"],
        ["TAX", "Tax Cases", "🟣"],
        ["PATENT CASE", "Patent Case", "🟣"],
        ["SLP", "Special Leave Petition", "🟣"],
        ["MAT. SUIT / MAT. REF.", "Matrimonial Suit / Reference", "🟢"],
        ["T.S.", "Title Suit", "🔵"],
        ["MONEY SUIT", "Money Suit", "🔵"],
        ["ADM. SUIT", "Admiralty Suit", "🔵"],
    ],
    widths=[1.9, 3.4, 0.7],
)

doc.add_page_break()

heading("District Court — Criminal", size=12, color=NAVY, space_before=8)
make_table(
    ["Code / name", "Flow"],
    [
        ["Criminal Case / Sessions Case / Summary Case", "🔴"],
        ["Cr. Appeal / Cr. Revision / Criminal Miscellaneous", "🔴"],
        ["Cr. Case Complaint (O) / (P)", "🔴"],
        ["Regular Bail (BP) / Anticipatory Bail (ABP) / Cri. Bail Appln.", "🔴"],
        ["NDPS Special Case / NDPS M.A. / NDPS Cri. Revn.", "🔴"],
        ["Spl. Case (POCSO / SC-ST / Excise / PMLA / PC Act / Vigilance …)", "🔴"],
        ["Indian Penal Code (IPC) / Harijan Atrocities Act / Excise Act", "🔴"],
        ["Domestic Violence Act 2005 / D.V. Appeal / D.V. Misc.", "🔴"],
        ["Electricity (E.C. Act) / Weights & Measures / Food Adulteration", "🔴"],
    ],
    widths=[5.3, 0.7],
)

heading("District Court — Family / Civil / Other", size=12, color=NAVY, space_before=10)
make_table(
    ["Code / name", "Flow"],
    [
        ["Matrimonial Case / Spl. Marriage Petition", "🟢"],
        ["Guardianship / Guardian & Wards Case", "🟢"],
        ["Maintenance", "🟢"],
        ["Title Suit / Title Partition / Title Mortgage / Title Eviction Suit", "🔵"],
        ["Money Suit / Money Appeal / Summary Suit / Small Cause Suit", "🔵"],
        ["Declaratory Suit / Injunction / Specific Performance / Eviction Suit", "🔵"],
        ["Succession / Succession Certificate / Probate / Letter of Administration", "🔵"],
        ["Execution / Darkhast / Final Decree", "🔵"],
        ["Arbitration Case / Caveat / Claim Case / Land Acquisition", "🔵"],
        ["Rent Suit / Rent Appeal / Trust Suit / Trust Appeal", "🔵"],
        ["Misc. Civil Appeal / Title Appeal / Munci. Appeal", "🟣"],
        ["Contempt Proceedings / Review Application / Insolvency", "🟣"],
        ["(100+ interlocutory applications: amendment, discovery, attachment, etc.)", "🔵"],
    ],
    widths=[5.3, 0.7],
)
para("")
para("Note: the District lists contain 100+ procedural / interlocutory application types "
     "(amendment of pleadings, discovery, attachment before judgment, substitution …). These are "
     "steps inside a case, not separate matters — they map to the Civil flow by default. Flag any "
     "you want treated as their own case type.", italic=True, color=GREY, size=10)
para("")
change_box("✍️  WHICH TYPES SHOULD SWITCH FLOWS?")

doc.save(OUT)
print("Saved:", OUT)
